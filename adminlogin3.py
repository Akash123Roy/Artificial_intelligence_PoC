import streamlit as st
import pandas as pd
import google.generativeai as genai
from io import BytesIO
from docx import Document
from docx.shared import Pt
import os
import json
from pandas import json_normalize
import re

# https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent?key=AIzaSyBjKe4Wk6CUtT0oSG1pUaq4Sn0ER90JpGY
# Configure Generative AI with API Key
genai.configure(api_key="")
model = genai.GenerativeModel(model_name="gemini-1.5-flash")


# Generate a summary based on the user prompt and input text
def generate_summary(text, prompt):
    full_prompt = f"{prompt}, {text}"
    response = model.generate_content(full_prompt)
    return response.text

# Create a DOC file from the generated summary text
def create_document(text):
    doc = Document()
    if isinstance(text, str):
        paragraphs = text.strip().split("\n")
    else:
        raise TypeError(f"The variable 'summary' is not a string; it's of type {type(summary)}")
    
    # Add paragraphs to the document
    for paragraph in paragraphs:
        if "**" in paragraph:
            parts = paragraph.split("**")
            p = doc.add_paragraph()
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(paragraph)
    
    # Set font size for readability
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

    # Save DOC to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Extract structured data from DOC for Excel conversion
def extract_data_from_doc(doc_file):
    doc = Document(doc_file)
    data = {}
    current_column = None

    # Parse DOC paragraphs to detect main points and descriptions
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Use numbering or bullet-like symbols to identify main points
        if text and (text[0].isdigit() or text.startswith('"')):
            current_column = text
            data[current_column] = []
        
        # Treat non-header text as associated data
        elif current_column and text:
            data[current_column].append(text)

    # Equalize row lengths across columns by padding with empty strings
    max_len = max((len(v) for v in data.values()), default=0)
    for key in data:
        data[key] += [''] * (max_len - len(data[key]))

    # Convert data dictionary to DataFrame
    df = pd.DataFrame(data)
    return df



def flatten_json_to_excel(json_data, excel_file_path, num_columns=10):
    """
    Converts a JSON object to an Excel file with the specified number of columns.
    Handles nested JSON data by flattening it.

    Parameters:
        json_data (dict or list): The JSON data to convert.
        excel_file_path (str): The full path and name of the output Excel file.
        num_columns (int): The maximum number of columns to include in the Excel.
    """
    # Flatten the JSON data using json_normalize
    if isinstance(json_data, list):
        flat_data = json_normalize(json_data)
    else:
        flat_data = json_normalize([json_data])
    
    # Truncate or limit columns to the specified number
    if len(flat_data.columns) > num_columns:
        flat_data = flat_data.iloc[:, :num_columns]
    
    # Create the directory if it doesn't exist
    output_dir = os.path.dirname(excel_file_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Save the data to an Excel file
    flat_data.to_excel(excel_file_path, index=False)
    print(f"Data successfully written to {excel_file_path}")

# def extract_data_from_doc(doc_file):
#     doc = Document(doc_file)
#     data = {}
#     current_column = None

#     # Parse document paragraphs
#     for para in doc.paragraphs:
#         text = para.text.strip()
        
#         # Heuristic for identifying main points (titles/headings)
#         if text and text.istitle():  # Assuming titles are written in title case
#             current_column = text
#             data[current_column] = []
        
#         # Treat other paragraphs as associated data under the current main point
#         elif current_column and text:
#             data[current_column].append(text)

#     # Padding each column to the same length for DataFrame conversion
#     max_len = max((len(v) for v in data.values()), default=0)
#     for key in data:
#         data[key] += [''] * (max_len - len(data[key]))

#     # Convert dictionary to DataFrame
#     df = pd.DataFrame(data)
#     return df


# Convert DataFrame to Excel format
def convert_to_excel(dataframe):
    excel_io = BytesIO()
    with pd.ExcelWriter(excel_io, engine='openpyxl') as writer:
        dataframe.to_excel(writer, index=False)
    excel_io.seek(0)
    return excel_io

# Streamlit app with login functionality and use-case selection
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

st.sidebar.header("Admin Login")

# Login handling
if not st.session_state.logged_in:
    username = st.sidebar.text_input("Username", placeholder="Enter your username")
    password = st.sidebar.text_input("Password", placeholder="Enter your password", type="password")
    if st.sidebar.button("Login"):
        if username == "admin" and password == "password123":
            st.sidebar.success("Login successful!")
            st.session_state.logged_in = True
        else:
            st.sidebar.error("Invalid username or password.")
else:
    st.sidebar.success("Logged in as admin.")

if st.session_state.logged_in:
    st.title("Welcome to LLM Usecase")
    st.sidebar.write("Model Selection")
    opt1 = st.sidebar.selectbox("Select LLM Model:", ("Select a Model", "Gemini-1.5-Flash"))
    st.sidebar.write("Use case selection")
    opt2 = st.sidebar.selectbox("Select use case:", ("Change Summary", "ARP Document Generation"))

    if opt2 == "Change Summary":
        opt3 = st.selectbox("Upload Type", ("Text", "Excel", "CSV"))
        if opt3 == "Excel":
            user_prompt = st.text_area("User Prompt:", height=100)
            upload_file = st.file_uploader("Upload Excel file here", type=["xlsx", "xls"])
            if upload_file is not None:
                xl_data = pd.ExcelFile(upload_file)
                sheet_names = xl_data.sheet_names
                sheet_name = st.selectbox("Select a Sheet", sheet_names)
                df = pd.read_excel(upload_file, sheet_name=sheet_name)
                st.write(f'Showing {sheet_name}:')
                df = df.astype(str)
                st.dataframe(df)

                
                

        elif opt3 == "Text":
            user_prompt = st.text_area("Enter Your Prompt here:", height=100)
            input_text = st.text_area("Paste your data here:")
            # input_tokens = model.count_tokens(input_text)
            
            if st.button("Generate Document"):
                if input_text:
                    total_input = f"{user_prompt},{input_text}"
                    input_tokens = model.count_tokens(total_input)
                    summary = generate_summary(input_text, user_prompt)
                    output_tokens = model.count_tokens(summary)
                    print(f"input tokens:{input_tokens} \n output tokens: {output_tokens}")
                    st.write(summary)
                    
                    # Generate DOC file
                    # doc_file = create_document(summary)
                    cleaned_response = summary.split('json',1)[1].strip()
                    cleaned_response = cleaned_response.split('}',1)[0].strip()+'}'
                    # print(f"cleaned resp: {cleaned_response}")
                    json_data = json.loads(cleaned_response)
                    # st.download_button(
                    #     label="Download DOC File",
                    #     data=doc_file,
                    #     file_name="summary.doc",
                    #     mime="application/msword"
                    # )
                    
                    # Convert DOC to structured Excel
                    # doc_file.seek(0)
                    # data = extract_data_from_doc(doc_file)

                    # excel_file = convert_to_excel(data)
                    df = pd.DataFrame([json_data])
                    excel_file = convert_to_excel(df)
                    # print(f"input tokens: {input_tokens}\n output tokens: {output_tokens}")
                    # Download Excel file
                    st.download_button(
                        label="Download Excel File",
                        data=excel_file,
                        file_name="summary.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                else:
                    st.warning("Please paste some text to generate summary")
