import streamlit as st
import pandas as pd
import google.generativeai as genai
from io import BytesIO
from docx import Document
from docx.shared import Pt


genai.configure(api_key="AIzaSyBjKe4Wk6CUtT0oSG1pUaq4Sn0ER90JpGY")
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

def generate_summary(text, prompt):
     full_prompt = f"{prompt},{text}"
     response = model.generate_content(full_prompt)
     return response.text

def create_document(text):
    doc = Document()  # Create a new Word document
    # Ensure that the text is a string before splitting
    if isinstance(text, str):
        paragraphs = text.strip().split("\n")
    else:
        raise TypeError(f"The variable 'summary' is not a string; it's of type {type(summary)}")
    for paragraph in paragraphs:
    # Replace ** with actual bold formatting
        if "**" in paragraph:
            parts = paragraph.split("**")
            p = doc.add_paragraph()  # Add paragraph to the Document
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    p.add_run(part).bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(paragraph)

    # Adjusting font size for readability
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)

                # Save the document to a BytesIO object
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio



# Initialize session state for login status
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

# Sidebar header for login form
st.sidebar.header("Admin Login")

# If not logged in, show login form
if not st.session_state.logged_in:
    # Username and password input fields
    username = st.sidebar.text_input("Username", placeholder="Enter your username")
    password = st.sidebar.text_input("Password", placeholder="Enter your password", type="password")

    # Login button
    if st.sidebar.button("Login"):
        # Check credentials (you can replace with your own validation logic)
        if username == "admin" and password == "password123":  # Replace with actual credentials
            st.sidebar.success("Login successful!")
            st.session_state.logged_in = True
        else:
            st.sidebar.error("Invalid username or password.")
else:
    # Once logged in, display the admin features
    st.sidebar.success("Logged in as admin.")


if st.session_state.logged_in:
     st.title("Welcome to LLM usecase")
     
     #creating menu
     st.sidebar.write("Model Selection")
     opt1 = st.sidebar.selectbox("Select LLM Model:", ("Select a Model", "Gemini-1.5-Flash"))
     st.sidebar.write("Use case selection")
     opt2 = st.sidebar.selectbox("Select use case:", ("Change Summary","ARP Document Generation"))

     if opt2 == "Change Summary":
          opt3 = st.selectbox("Upload Type", ("Text","Excel","CSV"))
          if opt3 == "Excel":
               user_prompt = st.text_area("User Prompt:", height=100)
               upload_file = st.file_uploader("Upload Excel file here", type = ["xlsx", "xls"])
               if upload_file is not None:
                    # read excel file
                    xl_data = pd.ExcelFile(upload_file)
                    sheet_names = xl_data.sheet_names

                    #let the user to select a sheet to display
                    sheet_name = st.selectbox("Select a Sheet", sheet_names)
                    # read the selected sheets into a dataframe
                    df = pd.read_excel(upload_file, sheet_name= sheet_name)
                    #display sheet
                    st.write(f'<div id="sel-head1>"Showing{sheet_name}: </div>', unsafe_allow_html= True)
                    # converting all data to string
                    df = df.astype(str)
                    st.dataframe(df, hide_index= True)
                
          elif opt3 == "Text":
               user_prompt = st.text_area("Enter Your Prompt here:", height= 100)
               input_text = st.text_area("Paste your data here:")
               if st.button("Generate Document"):
                    if input_text:
                         summary = generate_summary(user_prompt, input_text)
                         st.write(summary)
                         doc_file = create_document(summary)
                         st.download_button(
                              label = "Download File",
                              data = doc_file,
                              file_name= "summary.doc",
                              mime = "application/msword"
                         )
                    else:
                         st.warning("please paset some text to generate summary")


