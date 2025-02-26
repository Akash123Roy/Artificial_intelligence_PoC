import pandas as pd
import google.generativeai as genai
import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt

# Configure the API key
genai.configure(api_key="")
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

# Function to generate summary
def generate_summary(row):
    input_text = f"CR_NO: {row['CR_NO']}\n:Change Description {row['Change Description']}\nJustification: {row['Justification']}\nGenerate summary in 100 words."
    response = model.generate_content([input_text])
    return response.text

def generate_summary_from_text(text_file):
    prompt = f"Generate summary in 100 words{text_file}"
    response = model.generate_content(prompt)
    return response.text

# Streamlit UI
st.title("Summary Generator")

# Input option selection
option = st.selectbox(
    "Choose input type",
    ("Text", "Excel", "CSV")
)

if option == "Excel":
    file = st.file_uploader(f"Choose a {option} file", type=["xlsx"])
    if file:
        if st.button("Upload and Generate Summary"):
            buffer = BytesIO()
            df = pd.read_excel(file)  # Read the Excel file
            df['Summary'] = df.apply(generate_summary, axis=1)
            df.to_excel(buffer, index=False, engine='openpyxl')
            buffer.seek(0)
            st.write(df.head())
            st.download_button(
                label="Download Summary",
                data=buffer,
                file_name="summary.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

elif option == "CSV":
    file = st.file_uploader(f"Choose a {option} file", type=["csv"])
    if file:
        if st.button("Upload and Generate Summary"):
            buffer = BytesIO()
            df = pd.read_csv(file)  # Read the CSV file
            df['Summary'] = df.apply(generate_summary, axis=1)
            df.to_csv(buffer, index=False)
            buffer.seek(0)
            st.write(df.head())
            st.download_button(
                label="Download Summary",
                data=buffer,
                file_name="summary.csv",
                mime="text/csv"
            )

elif option == "Text":
    text_input = st.text_area("Paste your text here")
    if st.button("Generate Summary"):
        if text_input:
            summary = generate_summary_from_text(text_input)
            st.write("Summary:")
            st.write(summary)

            # Function to create a Word document
            def create_document():
                doc = Document()  # Create a new Word document

                # Ensure that the text is a string before splitting
                if isinstance(summary, str):
                    paragraphs = summary.strip().split("\n")
                else:
                    raise TypeError(f"The variable 'summary' is not a string; it's of type {type(summary)}")

                for paragraph in paragraphs:
                    # Replace ** with actual bold formatting
                    if "**" in paragraph:
                        parts = paragraph.split("**")
                        p = doc.add_paragraph()  # Add paragraph to the Document
                        for i, part in enumerate(parts):
                            if i % 2 == 1:
                                p.add_run(part).bold = True
                            else:
                                p.add_run(part)
                    else:
                        doc.add_paragraph(paragraph)

                # Adjusting font size for readability
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

                # Save the document to a BytesIO object
                bio = BytesIO()
                doc.save(bio)
                bio.seek(0)

                return bio

            # Separate button for generating and downloading the document
            doc_file = create_document()
            st.download_button(
                label="Download Word Document",
                data=doc_file,
                file_name="formatted_document.docx",
                mime="application/msword"
            )

        else:
            st.warning("Please paste some text to generate a summary.")
